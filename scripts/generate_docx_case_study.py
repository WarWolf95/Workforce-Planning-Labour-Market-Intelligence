"""
Generates a professional Word Document (.docx) Case Study / Project Briefing
suitable for recruiters, hiring managers, and portfolio review.
"""

from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from utils import setup_logging

logger = setup_logging("generate_docx_case_study")

WORKSPACE_DIR = Path(__file__).resolve().parents[1]
REPORTS_DIR = WORKSPACE_DIR / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
DOCX_FILE = REPORTS_DIR / "Workforce_Planning_Case_Study.docx"

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def main() -> None:
    logger.info("=========================================")
    logger.info("STARTING CASE STUDY DOCX GENERATION")
    logger.info("=========================================")
    
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles config
    styles = doc.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Charcoal
    
    # ----------------------------------------------------
    # TITLE & HEADER (No title page needed for clean briefing)
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Workforce Planning & Labour Market Intelligence")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark Blue
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("Project Portfolio Case Study & Executive Briefing")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)  # Gray
    
    # Add metadata block
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.autofit = False
    
    col_widths = [Inches(3.25), Inches(3.25)]
    for i, row in enumerate(meta_table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            
    # Set labels
    meta_table.cell(0, 0).paragraphs[0].add_run("AUTHOR: Maulik Parmar").bold = True
    meta_table.cell(0, 1).paragraphs[0].add_run("TARGET AUDIENCE: C-Suite / Portfolio Reviewers").bold = True
    meta_table.cell(1, 0).paragraphs[0].add_run("VERSION: 1.0 (Production Grade)").bold = True
    meta_table.cell(1, 1).paragraphs[0].add_run("DATE: July 2026").bold = True
    
    # Formatting spacing of metadata table
    for row in meta_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)
            set_cell_background(cell, "F2F2F2")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE BRIEFING
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    r1 = h1.add_run("1. Executive Briefing")
    r1.font.name = "Arial"
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    p = doc.add_paragraph(
        "This project models and benchmarks our internal corporate workforce roster against external "
        "macroeconomic indexes to identify retention, succession, and salary mismatch vulnerabilities. "
        "The analysis integrates ONS ASHE (Annual Survey of Hours and Earnings), ONS Labour Force Survey, "
        "and active scraped vacancy listings (Adzuna) across three key UK shortage sectors: "
        "Clinical Operations, Digital Services, and Green Energy Projects."
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Sector findings
    p_lead = doc.add_paragraph("Key Risk Vectors Identified:")
    p_lead.paragraph_format.space_after = Pt(4)
    
    bullets = [
        ("Digital Services", "Comp deficit of -16.2% on average against ONS medians. Software Developers in London lag the market by -26.3% (-£17,133 deficit), driving extreme turnover risk."),
        ("Green Energy Projects", "Demographic hazard with 34.7% to 41.0% of engineers aged 55+. Critical lack of succession planning—Battery Design Engineers have a 100% succession gap for critical over-50 staff."),
        ("Clinical Operations", "Nursing and Support structures face intense vacancy pressure. Staff Nurse posts have a 26.4% retirement risk coupled with 14 positions lacking designated successors.")
    ]
    
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Inches(0.25)
        run_bold = bp.add_run(f"{b_title}: ")
        run_bold.bold = True
        bp.add_run(b_desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    r2 = h2.add_run("2. Technical & Data Architecture")
    r2.font.name = "Arial"
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    p = doc.add_paragraph(
        "The portfolio implements a highly scalable, production-grade data pipeline designed for local "
        "testing and enterprise execution:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    tech_stack = [
        ("Orchestrator (run_pipeline.py)", "Coordinates modular components, executing data steps in chronological sequence and exporting analytical reports directly from raw queries."),
        ("Text Processing & TF-IDF (scikit-learn)", "Extracts and cleans technical skills keywords from unstructured corporate job descriptions, mapping capabilities against our SOC taxonomy."),
        ("Analytical Cache (DuckDB)", "Stages raw ONS, internal, and scraped CSV records into memory and compiles views for complex gap modeling."),
        ("Enterprise DB Compilation (Oracle SQL & SQLite)", "Compiles tables, schemas, and checks. It downsamples vacancy data to 600 records to allow seamless deployment inside Oracle LiveSQL limits.")
    ]
    
    for tech, desc in tech_stack:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.left_indent = Inches(0.25)
        run_bold = bp.add_run(f"{tech}: ")
        run_bold.bold = True
        bp.add_run(desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # SECTION 3: KEY KPI VERIFICATION
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    r3 = h3.add_run("3. Strategic KPI Verification")
    r3.font.name = "Arial"
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    p = doc.add_paragraph(
        "To ensure mathematical consistency across reporting environments (Power BI, SQLite, DuckDB, and Oracle), "
        "the database assertion checks enforce the following specific KPI targets:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    # Add KPI Table
    kpi_table = doc.add_table(rows=3, cols=3)
    kpi_table.autofit = False
    
    headers = ["Role Profile", "Key Analytics Metric Checked", "Target Pipeline Value"]
    kpi_widths = [Inches(2.5), Inches(2.75), Inches(1.25)]
    
    hdr_cells = kpi_table.rows[0].cells
    for j, h_name in enumerate(headers):
        hdr_cells[j].text = h_name
        hdr_cells[j].width = kpi_widths[j]
        set_cell_background(hdr_cells[j], "1F497D")
        hdr_cells[j].paragraphs[0].runs[0].font.bold = True
        hdr_cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    kpi_data = [
        ("Software Developer (London)", "Avg Salary Lag vs ONS Median", "-26.3% (£47.9k vs £65.0k ONS)"),
        ("Battery Design Engineer", "Retirement Risk / Succession Gap", "41.0% Risk / 11 Roles No Successor")
    ]
    
    for i, row_data in enumerate(kpi_data):
        row_cells = kpi_table.rows[i+1].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val
            row_cells[j].width = kpi_widths[j]
            row_cells[j].paragraphs[0].paragraph_format.space_before = Pt(4)
            row_cells[j].paragraphs[0].paragraph_format.space_after = Pt(4)
            if i % 2 == 1:
                set_cell_background(row_cells[j], "F9FBFD")
            else:
                set_cell_background(row_cells[j], "FFFFFF")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # SECTION 4: RECOMMENDATIONS
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    r4 = h4.add_run("4. C-Suite Strategic Recommendations")
    r4.font.name = "Arial"
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    recs = [
        ("Compensation Realignment", "Implement a target 12% to 15% salary uplift for critical Software and Data Engineers in London and South East. This incurs a £180k adjustment cost but prevents talent flight, avoiding £350k in recruitment friction and contracting premiums (Net ROI: £170k)."),
        ("Knowledge Transfer Partnerships", "Establish a Graduate Apprenticeship mentorship pipeline in battery chemistry and grid connections. Pairing junior hires with retiring Battery Design Engineers protects IP and prevents renewable project delays."),
        ("Clinical Pipeline Consolidation", "Partner with local nursing colleges for clinical student placements, and review retention benefits (scheduling flexibility, career paths) to reduce staff nurse turnover, lowering agency costs by £220k.")
    ]
    
    for title, desc in recs:
        p_rec = doc.add_paragraph()
        p_rec.paragraph_format.space_after = Pt(4)
        run_bold = p_rec.add_run(f"Recommendation: {title}\n")
        run_bold.bold = True
        run_bold.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p_rec.add_run(desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Save the document
    doc.save(str(DOCX_FILE))
    logger.info(f"Successfully generated Case Study Word Document at: {DOCX_FILE}")
    logger.info("=========================================")
    logger.info("CASE STUDY DOCX GENERATION COMPLETED")
    logger.info("=========================================")

if __name__ == "__main__":
    main()
